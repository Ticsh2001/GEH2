import os
import json
import re
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED

import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


PARAM_EXPORT_COLUMNS = ["KKS код", "Описание", "Ед. изм.", "Используемые сигналы", "Код"]
RULE_EXPORT_COLUMNS = ["KKS код", "Используемые сигналы", "Код", "Описание", "Методические указания"]
RULE_EXPORT_EXTRA_COLUMNS = ["params_list", "rule_code_raw", "possibleCause"]


def _validate_filename(filename: str) -> bool:
    if not isinstance(filename, str) or not filename:
        return False
    if ".." in filename or "/" in filename or "\\" in filename:
        return False
    return filename.endswith(".json")


def _iter_elements(elements):
    if isinstance(elements, dict):
        return elements.values()
    if isinstance(elements, list):
        return elements
    return []


def _unique_preserve_order(items):
    seen = set()
    out = []
    for x in items:
        if not x:
            continue
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out


def prepare_code_for_system(code_str: str, input_signal_names=None) -> str:
    """
    Python-версия JS-функции prepareCodeForSystem с поддержкой двух режимов.
    """
    if not code_str or not isinstance(code_str, str):
        return code_str

    # Проверяем наличие термодинамических функций
    thermo_map = {
        "ENTHALPY_PS": "enthalpy_ps",
        "ENTHALPY_PT": "enthalpy_pt",
        "PRESSURE_SATURATION": "pressure_saturation",
        "TEMPERATURE_SATURATION": "temperature_saturation",
        "ENTROPY_PT": "entropy_pt",
        "TEMPERATURE_PS": "temperature_ps"
    }
    has_thermo = any(fn + '(' in code_str for fn in thermo_map)

    if not has_thermo:
        # ===== СТАРЫЙ РЕЖИМ (без изменений) =====
        out = code_str

        out = re.sub(r"\bAND\b", "&&", out)
        out = re.sub(r"\bOR\b", "||", out)
        out = re.sub(r"\bNOT\b", "!", out)

        out = out.replace("§", "_")

        out = re.sub(r"(?<![<>=!])=(?![=])", "==", out)

        try:
            used_signals = []
            for name in (input_signal_names or []):
                if name is None:
                    continue
                name = str(name).strip()
                if not name:
                    continue
                name = name.replace("§", "_")
                used_signals.append(name)

            unique_signals = _unique_preserve_order(used_signals)
            starts_with_digit = [name for name in unique_signals if re.match(r"^\d", name)]

            for sig in starts_with_digit:
                pattern = re.compile(rf"(^|[^\w.])({re.escape(sig)})(?![\w.])")
                out = pattern.sub(r"\1P\2", out)
        except Exception as e:
            print(f"[WARN] prepare_code_for_system: {e}")

        fn_list = [
            "PREV", "GETPOINT", "INTERPOLATE",
            "HISTORYAVG", "HISTORYCOUNT", "HISTORYSUM",
            "HISTORYMAX", "HISTORYMIN", "HISTORYDIFF", "HISTORYGRADIENT"
        ]
        for fn in fn_list:
            pattern = re.compile(rf"\b{fn}\s*\(\s*([^,\)]+)")
            def repl(match):
                p1 = match.group(1)
                if re.match(r"^['\"]", p1.strip()):
                    return match.group(0)
                arg = p1.strip()
                arg = re.sub(r"^P(?=\d)", "", arg)
                return f"{fn}('{arg}'"
            out = pattern.sub(repl, out)

        out = re.sub(r"(^|[(+*/%,\s!-]|==|!=|<=|>=|<|>|&&|\|\|)-(?=P\d)", r"\1-1 * ", out)
        out = re.sub(r"(^|[(+*/%,\s!-]|==|!=|<=|>=|<|>|&&|\|\|)-(?=\()", r"\1-1 * ", out)

        return out

    # ===== НОВЫЙ РЕЖИМ =====
    out = code_str

    # 1. Логические операторы → строчные слова
    out = re.sub(r"\bAND\b", "and", out)
    out = re.sub(r"\bOR\b", "or", out)
    out = re.sub(r"\bNOT\b", "not", out)

    # 2. § → _
    out = out.replace("§", "_")

    # 3. Одиночное '=' → '=='
    out = re.sub(r"(?<![<>=!])=(?![=])", "==", out)

    # 4. Переводим имена функций в нижний регистр
    function_name_map = {
        "WHEN": "when",
        "ABS": "abs",
        "EXP": "exp",
        "POW": "pow",
        "LOG": "log",
        "LOG10": "log10",
        "MIN": "min",
        "MAX": "max",
        "AVG": "avg",
        "MED": "med",
        "ROUND": "round",
        "VARIANCE": "variance",
        "STDEV": "stdev",
        "HISTORYAVG": "history_avg",
        "HISTORYMIN": "history_min",
        "HISTORYMAX": "history_max",
        "HISTORYDIFF": "history_diff",
        "HISTORYDIFFMAX": "history_diff_max",
        "PREV": "prev",
        "GETPOINT": "getpoint",
        "INTERPOLATE": "interpolate",
        "ENTHALPY_PS": "enthalpy_ps",
        "ENTHALPY_PT": "enthalpy_pt",
        "PRESSURE_SATURATION": "pressure_saturation",
        "TEMPERATURE_SATURATION": "temperature_saturation",
        "ENTROPY_PT": "entropy_pt",
        "TEMPERATURE_PS": "temperature_ps"
    }

    for upper, lower in function_name_map.items():
        # Замена только перед открывающей скобкой
        out = re.sub(r'\b' + re.escape(upper) + r'\s*\(', lower + '(', out)

    # 5. Оборачиваем оставшиеся идентификаторы в фигурные скобки
    known_functions = set(function_name_map.values()) | {"and", "or", "not"}

    def wrap_identifier(match):
        token = match.group(0)
        if re.match(r"^\d+(\.\d+)?(e[+-]?\d+)?$", token, re.IGNORECASE):
            return token
        if token.lower() in known_functions:
            return token
        return "{" + token + "}"

    out = re.sub(
        r"(?<![A-Za-z0-9_§.])[A-Za-z0-9_§]+(?![A-Za-z0-9_§.])",
        wrap_identifier,
        out
    )

    # 6. Унарный минус в новом режиме не обрабатываем

    return out


def _load_project_payloads(filenames: list[str], project_dir: str):
    payloads = []
    missing = []

    seen = set()
    for filename in filenames:
        if filename in seen:
            continue
        seen.add(filename)

        if not _validate_filename(filename):
            missing.append(filename)
            continue

        path = os.path.join(project_dir, filename)
        if not os.path.exists(path):
            missing.append(filename)
            continue

        try:
            with open(path, "r", encoding="utf-8") as f:
                payloads.append((filename, json.load(f)))
        except Exception:
            missing.append(filename)

    return payloads, missing


def _build_parameter_rows(project_payloads):
    rows = []

    for _, data in project_payloads:
        project = data.get("project") or {}
        if project.get("type") != "parameter":
            continue

        input_signals = []

        for el in _iter_elements(data.get("elements") or {}):
            if not isinstance(el, dict):
                continue
            if el.get("type") != "input-signal":
                continue

            props = el.get("props") or {}
            name = props.get("name")

            if isinstance(name, str):
                name = name.strip()
            elif name is not None:
                name = str(name)
            else:
                name = ""

            input_signals.append(name)

        input_signals = _unique_preserve_order(input_signals)

        rows.append({
            "KKS код": project.get("code", ""),
            "Описание": project.get("description", ""),
            "Ед. изм.": project.get("dimension", ""),
            "Используемые сигналы": "; ".join(input_signals),
            "Код": prepare_code_for_system(data.get("code", ""), input_signals),
        })

    return rows


def _build_rule_rows(project_payloads):
    rows = []

    for _, data in project_payloads:
        project = data.get("project") or {}
        if project.get("type") != "rule":
            continue

        display_signals = []
        input_signal_names = []
        params_list = []

        for el in _iter_elements(data.get("elements") or {}):
            if not isinstance(el, dict):
                continue

            el_type = el.get("type")
            if el_type not in ("input-signal", "table"):
                continue

            props = el.get("props") or {}
            name = props.get("name") or ""
            desc = props.get("description") or ""

            if isinstance(name, str):
                name = name.strip()
            else:
                name = str(name)

            params_list.append((name, desc))
            display_signals.append(name)

            if el_type == "input-signal":
                input_signal_names.append(name)

        display_signals = _unique_preserve_order(display_signals)
        input_signal_names = _unique_preserve_order(input_signal_names)

        rows.append({
            "KKS код": project.get("code", ""),
            "Используемые сигналы": "; ".join(display_signals),
            "Код": prepare_code_for_system(data.get("code", ""), input_signal_names),
            "Описание": project.get("description", ""),
            "Методические указания": project.get("guidelines", ""),
            "params_list": params_list,
            "rule_code_raw": project.get("code", ""),
            "possibleCause": project.get("possibleCause", ""),   # ← возможная причина
        })

    return rows


def _build_parameters_xlsx_bytes(df: pd.DataFrame) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)

    buffer.seek(0)
    wb = load_workbook(buffer)
    ws = wb.active

    wrap_cols = {"Описание", "Используемые сигналы", "Код"}
    header = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    col_idx = {name: i + 1 for i, name in enumerate(header)}

    for col_name in wrap_cols:
        idx = col_idx.get(col_name)
        if not idx:
            continue
        for row in ws.iter_rows(min_row=2, min_col=idx, max_col=idx):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    widths = {
        "KKS код": 18,
        "Описание": 50,
        "Ед. изм.": 12,
        "Используемые сигналы": 30,
        "Код": 60,
    }

    for name, width in widths.items():
        idx = col_idx.get(name)
        if idx:
            ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = width

    out = BytesIO()
    wb.save(out)
    return out.getvalue()

def _build_parameters_docx_bytes(df: pd.DataFrame) -> bytes:
    """Генерирует DOCX-файл с параметрами, каждый параметр на отдельной странице."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    def normalize_text(val):
        if val is None:
            return ""
        try:
            if isinstance(val, float) and pd.isna(val):
                return ""
        except Exception:
            pass
        s = str(val)
        return s.replace("\r\n", "\n").replace("\r", "\n").replace("\\n", "\n")

    def add_paragraph_line(text="", bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.5

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        return p

    def add_multiline(text="", bold=False):
        lines = normalize_text(text).split("\n")
        if not lines:
            add_paragraph_line("", bold=bold)
            return
        for line in lines:
            add_paragraph_line(line, bold=bold)

    # Поля из PARAM_EXPORT_COLUMNS
    fields = [
        ("KKS код", "KKS код"),
        ("Описание", "Описание"),
        ("Ед. изм.", "Ед. изм."),
        ("Используемые сигналы", "Используемые сигналы"),
        ("Код", "Код"),
    ]

    for idx, row in df.iterrows():
        # Первый параметр — без номера, если нужно, можно добавить "Параметр N"
        for label, col in fields:
            add_paragraph_line(f"{label}:", bold=True)
            add_multiline(normalize_text(row.get(col, "")))

        if idx != len(df) - 1:
            doc.add_page_break()

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_rules_docx_bytes(df: pd.DataFrame) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    def normalize_text(val):
        if val is None:
            return ""
        try:
            if isinstance(val, float) and pd.isna(val):
                return ""
        except Exception:
            pass
        s = str(val)
        return s.replace("\r\n", "\n").replace("\r", "\n").replace("\\n", "\n")

    def add_paragraph_line(text="", bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.5

        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        return p

    def add_multiline(text="", bold=False):
        lines = normalize_text(text).split("\n")
        if not lines:
            add_paragraph_line("", bold=bold)
            return
        for line in lines:
            add_paragraph_line(line, bold=bold)

    for idx, row in df.iterrows():
        n = idx + 1
        rule_code = normalize_text(row.get("rule_code_raw"))
        code_text = normalize_text(row.get("Код"))
        descr = normalize_text(row.get("Описание"))
        guidelines = normalize_text(row.get("Методические указания"))
        possible_cause = normalize_text(row.get("possibleCause"))
        params_list = row.get("params_list") or []

        add_paragraph_line(f"Правило {n}: ({rule_code})", bold=True)

        add_paragraph_line("Краткое описание:", bold=True)
        add_multiline(descr)

        add_paragraph_line("Причина:", bold=True)
        add_multiline(possible_cause)

        add_paragraph_line("Логика определения:", bold=True)
        add_paragraph_line("")

        add_paragraph_line("Формула:", bold=True)
        add_multiline(code_text)

        add_paragraph_line("Используемые параметры:", bold=True)
        if params_list:
            for name, desc in params_list:
                name_t = normalize_text(name)
                desc_t = normalize_text(desc).replace("\n", "; ")
                line = f"- {name_t}"
                if desc_t:
                    line += f" — {desc_t}"
                add_paragraph_line(line)
        else:
            add_paragraph_line("- (нет)")

        add_paragraph_line("Время срабатывания:", bold=True)
        add_paragraph_line("")

        add_paragraph_line("Методические указания:", bold=True)
        add_multiline(guidelines)

        if idx != len(df) - 1:
            doc.add_page_break()

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def _build_combined_excel_bytes(rows: list[dict]) -> bytes:
    """Создаёт Excel-файл с объединёнными параметрами и правилами."""
    df = pd.DataFrame(rows, columns=PARAM_EXPORT_COLUMNS)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)
    buffer.seek(0)
    wb = load_workbook(buffer)
    ws = wb.active

    wrap_cols = {"Описание", "Используемые сигналы", "Код"}
    header = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    col_idx = {name: i + 1 for i, name in enumerate(header)}
    for col_name in wrap_cols:
        idx = col_idx.get(col_name)
        if not idx:
            continue
        for row in ws.iter_rows(min_row=2, min_col=idx, max_col=idx):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
    widths = {"KKS код": 18, "Описание": 50, "Ед. изм.": 12, "Используемые сигналы": 30, "Код": 60}
    for name, width in widths.items():
        idx = col_idx.get(name)
        if idx:
            ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = width

    out = BytesIO()
    wb.save(out)
    return out.getvalue()


def export_selected_projects(filenames: list[str], project_dir: str, param_format: str = "excel", export_rules_excel: bool = True) -> dict:
    if not isinstance(filenames, list) or not filenames:
        raise ValueError("Нужно передать непустой список filenames")

    if not project_dir or not os.path.isdir(project_dir):
        raise FileNotFoundError("Папка проектов не найдена")

    project_payloads, missing = _load_project_payloads(filenames, project_dir)

    if not project_payloads:
        raise FileNotFoundError("Не найдено ни одного проекта для экспорта")

    parameter_rows = _build_parameter_rows(project_payloads)
    rule_rows = _build_rule_rows(project_payloads)

    files_to_send = []

    if parameter_rows:
        df_param = pd.DataFrame.from_records(parameter_rows).reindex(columns=PARAM_EXPORT_COLUMNS)
        if param_format == "word":
            files_to_send.append({
                "filename": "Сигналы.docx",
                "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": _build_parameters_docx_bytes(df_param)
            })
        else:
            files_to_send.append({
                "filename": "Сигналы.xlsx",
                "media_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "content": _build_parameters_xlsx_bytes(df_param)
            })

    if rule_rows:
        df_rule = pd.DataFrame.from_records(rule_rows)
        for col in RULE_EXPORT_COLUMNS + RULE_EXPORT_EXTRA_COLUMNS:
            if col not in df_rule.columns:
                df_rule[col] = ""
        df_rule = df_rule[RULE_EXPORT_COLUMNS + RULE_EXPORT_EXTRA_COLUMNS]

        files_to_send.append({
            "filename": "Правила.docx",
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content": _build_rules_docx_bytes(df_rule)
        })
        if export_rules_excel:
            rule_excel_rows = []
            for r in rule_rows:
                rule_excel_rows.append({
                    "KKS код": r.get("KKS код", ""),
                    "Описание": r.get("Описание", ""),   # описание правила
                    "Ед. изм.": "",
                    "Используемые сигналы": r.get("Используемые сигналы", ""),
                    "Код": r.get("Код", "")
                })
            if rule_excel_rows:
                files_to_send.append({
                    "filename": "Правила.xlsx",
                    "media_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "content": _build_combined_excel_bytes(rule_excel_rows)  # используем ту же функцию форматирования
                })

    if not files_to_send:
        raise ValueError("Среди выбранных проектов нет типов parameter или rule")

    if len(files_to_send) == 1:
        return files_to_send[0]

    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, "w", ZIP_DEFLATED) as zf:
        for item in files_to_send:
            zf.writestr(item["filename"], item["content"])

        if missing:
            zf.writestr(
                "skipped.txt",
                "Не удалось прочитать следующие файлы:\n" + "\n".join(missing)
            )

    return {
        "filename": "Экспорт_проектов.zip",
        "media_type": "application/zip",
        "content": zip_buffer.getvalue()
    }

def get_code_length(payload: dict) -> int:
    """
    Возвращает длину обработанного кода проекта (после prepare_code_for_system),
    как это будет при экспорте.
    """
    project_meta = payload.get("project") or {}
    ptype = project_meta.get("type", "")
    raw_code = payload.get("code", "") or ""

    if ptype == "parameter":
        # Собираем input_signals как в _build_parameter_rows
        input_signals = []
        for el in _iter_elements(payload.get("elements") or {}):
            if isinstance(el, dict) and el.get("type") == "input-signal":
                name = (el.get("props") or {}).get("name")
                if isinstance(name, str):
                    name = name.strip()
                elif name is not None:
                    name = str(name)
                else:
                    name = ""
                input_signals.append(name)
        input_signals = _unique_preserve_order(input_signals)
        processed = prepare_code_for_system(raw_code, input_signals)
    elif ptype == "rule":
        # Собираем input_signal_names как в _build_rule_rows (только input-signal)
        input_signal_names = []
        for el in _iter_elements(payload.get("elements") or {}):
            if isinstance(el, dict) and el.get("type") in ("input-signal", "table"):
                name = (el.get("props") or {}).get("name", "")
                if isinstance(name, str):
                    name = name.strip()
                else:
                    name = str(name)
                if el.get("type") == "input-signal":
                    input_signal_names.append(name)
        input_signal_names = _unique_preserve_order(input_signal_names)
        processed = prepare_code_for_system(raw_code, input_signal_names)
    else:
        # Для других типов просто сырой код (или можно возвращать 0)
        processed = raw_code

    return len(processed)